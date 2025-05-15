"""
Sample business document generator for testing the AI Agent Workflow.
Creates a realistic business document with sections, tables, and charts.
"""
import os
import random
from datetime import datetime, timedelta
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from faker import Faker

# Initialize Faker
fake = Faker()

class SampleDocGenerator:
    """
    Generate a realistic business document with sections, tables, and charts.
    """
    
    def __init__(self, company_name=None, output_dir="samples"):
        """
        Initialize the generator.
        
        Args:
            company_name: Company name (random if None)
            output_dir: Directory to save the document
        """
        self.company_name = company_name or f"{fake.company()} {random.choice(['Inc.', 'Corp.', 'Ltd.', 'LLC'])}"
        self.output_dir = output_dir
        self.current_year = datetime.now().year
        self.doc = Document()
        
        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Configure document properties
        self._configure_document()
        
        # Generate financial data
        self.financial_data = self._generate_financial_data()
        
        # Generate other data
        self.strategic_goals = self._generate_strategic_goals()
        self.risks = self._generate_risks()
        self.competitors = self._generate_competitors()
    
    def _configure_document(self):
        """
        Configure document properties and styles.
        """
        # Set document properties
        properties = self.doc.core_properties
        properties.title = f"{self.company_name} Annual Business Plan"
        properties.author = f"{fake.name()}"
        properties.subject = "Business Plan"
        properties.keywords = "strategic plan, business, annual report"
        
        # Configure styles
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def _generate_financial_data(self):
        """
        Generate realistic financial data.
        
        Returns:
            Dictionary with financial data
        """
        base_revenue = random.randint(5000000, 100000000)
        growth_rate = random.uniform(0.05, 0.25)
        
        # Generate quarterly data for past year
        data = {
            "revenue": {
                "Q1": int(base_revenue * random.uniform(0.9, 1.1)),
                "Q2": int(base_revenue * random.uniform(0.9, 1.1)),
                "Q3": int(base_revenue * random.uniform(0.9, 1.1)),
                "Q4": int(base_revenue * random.uniform(0.9, 1.1))
            },
            "expenses": {},
            "profit": {},
            "growth_rate": growth_rate,
            "market_share": random.uniform(0.05, 0.3),
            "customer_retention": random.uniform(0.75, 0.95)
        }
        
        # Calculate expenses (60-80% of revenue)
        expense_ratio = random.uniform(0.6, 0.8)
        for quarter in ["Q1", "Q2", "Q3", "Q4"]:
            data["expenses"][quarter] = int(data["revenue"][quarter] * expense_ratio)
            data["profit"][quarter] = data["revenue"][quarter] - data["expenses"][quarter]
        
        # Calculate totals
        data["revenue"]["Total"] = sum(data["revenue"][q] for q in ["Q1", "Q2", "Q3", "Q4"])
        data["expenses"]["Total"] = sum(data["expenses"][q] for q in ["Q1", "Q2", "Q3", "Q4"])
        data["profit"]["Total"] = sum(data["profit"][q] for q in ["Q1", "Q2", "Q3", "Q4"])
        
        # Calculate projections for next year
        data["projections"] = {
            "revenue": int(data["revenue"]["Total"] * (1 + growth_rate)),
            "expenses": int(data["expenses"]["Total"] * (1 + growth_rate * 0.8)),
            "profit": int(data["profit"]["Total"] * (1 + growth_rate * 1.2))
        }
        
        return data
    
    def _generate_strategic_goals(self):
        """
        Generate strategic goals.
        
        Returns:
            List of strategic goals
        """
        growth_target = int(self.financial_data["growth_rate"] * 100)
        
        goals = [
            f"Achieve {growth_target}-{growth_target+5}% revenue growth by expanding {fake.bs()}",
            f"Enter {random.randint(2, 5)} new markets in the {random.choice(['Asia-Pacific', 'European', 'Latin American', 'North American'])} region",
            f"Launch {random.randint(1, 3)} new products in the {fake.catch_phrase()} category",
            f"Improve operational efficiency by {random.randint(5, 15)}% through process automation",
            f"Increase market share from {self.financial_data['market_share']*100:.1f}% to {self.financial_data['market_share']*100+random.randint(3, 8):.1f}%",
            f"Expand the {random.choice(['sales', 'engineering', 'customer support'])} team by {random.randint(10, 50)} new hires",
            f"Reduce carbon footprint by {random.randint(10, 30)}% through sustainable practices",
            f"Achieve customer satisfaction score of {random.randint(85, 98)}%"
        ]
        
        return random.sample(goals, k=min(6, len(goals)))
    
    def _generate_risks(self):
        """
        Generate business risks.
        
        Returns:
            List of risk items
        """
        risks = [
            {"category": "Competition", "description": f"Increasing pressure from {fake.company()} and other competitors offering similar solutions at lower price points"},
            {"category": "Technology", "description": f"Rapid advancements in {fake.catch_phrase()} may render current products obsolete"},
            {"category": "Market", "description": f"Market saturation in {random.choice(['domestic', 'primary', 'current'])} markets limiting organic growth potential"},
            {"category": "Regulatory", "description": f"New {random.choice(['data privacy', 'security', 'compliance', 'tax'])} regulations in key markets increasing operational overhead"},
            {"category": "Talent", "description": f"Difficulty recruiting qualified {random.choice(['engineers', 'data scientists', 'product managers'])} in competitive job market"},
            {"category": "Financial", "description": f"Currency fluctuations affecting profit margins in international markets"},
            {"category": "Supply Chain", "description": f"Dependencies on key suppliers creating potential bottlenecks in {fake.bs()}"},
            {"category": "Economic", "description": f"Economic uncertainty due to {random.choice(['inflation', 'recession concerns', 'interest rate changes'])}"}
        ]
        
        return random.sample(risks, k=min(6, len(risks)))
    
    def _generate_competitors(self):
        """
        Generate competitor information.
        
        Returns:
            List of competitor data
        """
        competitors = []
        market_share_left = 1.0 - self.financial_data["market_share"]
        
        for i in range(random.randint(3, 6)):
            if i < len(competitors) - 1:
                share = market_share_left * random.uniform(0.1, 0.4)
                market_share_left -= share
            else:
                share = market_share_left
            
            competitors.append({
                "name": fake.company(),
                "market_share": share,
                "strengths": random.sample([
                    "Strong brand recognition",
                    "Innovative product features",
                    "Aggressive pricing",
                    "Excellent customer support",
                    "Large customer base",
                    "Strong international presence",
                    "Proprietary technology",
                    "Vertical market expertise"
                ], k=random.randint(1, 3)),
                "weaknesses": random.sample([
                    "Limited product portfolio",
                    "Outdated technology",
                    "Poor customer service",
                    "Limited market reach",
                    "High employee turnover",
                    "Financial instability",
                    "Regulatory compliance issues",
                    "Slow release cycles"
                ], k=random.randint(1, 3))
            })
        
        return competitors
    
    def generate_document(self):
        """
        Generate the complete document.
        
        Returns:
            Path to generated document
        """
        self._add_cover_page()
        self._add_table_of_contents()
        self._add_executive_summary()
        self._add_market_analysis()
        self._add_competitor_analysis()
        self._add_strategic_goals()
        self._add_financial_projections()
        self._add_risk_assessment()
        self._add_implementation_plan()
        self._add_appendices()
        
        # Save the document
        docx_path = os.path.join(self.output_dir, f"{self.company_name.replace(' ', '_')}_Business_Plan.docx")
        self.doc.save(docx_path)
        
        # Convert to PDF (if possible)
        pdf_path = self._convert_to_pdf(docx_path)
        
        return pdf_path or docx_path
    
    def _add_cover_page(self):
        """
        Add a cover page to the document.
        """
        # Add title
        self.doc.add_paragraph().add_run().add_break()
        title = self.doc.add_paragraph()
        title.alignment = 1  # Center
        title_run = title.add_run(f"{self.company_name}\nAnnual Business Plan\n{self.current_year}")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # Add subtitle
        self.doc.add_paragraph().add_run().add_break()
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = 1  # Center
        subtitle.add_run(f"Prepared by: Executive Leadership Team\nDate: {fake.date_this_year().strftime('%B %d, %Y')}").font.size = Pt(14)
        
        # Add confidentiality notice
        self.doc.add_paragraph().add_run().add_break()
        notice = self.doc.add_paragraph()
        notice.alignment = 1  # Center
        notice.add_run("CONFIDENTIAL").font.bold = True
        
        # Add page break
        self.doc.add_page_break()
    
    def _add_table_of_contents(self):
        """
        Add a table of contents to the document.
        """
        self.doc.add_heading("Table of Contents", level=1)
        toc = self.doc.add_paragraph()
        sections = [
            "1. Executive Summary",
            "2. Market Analysis",
            "3. Competitor Analysis",
            "4. Strategic Goals",
            "5. Financial Projections",
            "6. Risk Assessment",
            "7. Implementation Plan",
            "8. Appendices"
        ]
        
        for section in sections:
            toc.add_run(f"{section}\n")
        
        self.doc.add_page_break()
    
    def _add_executive_summary(self):
        """
        Add executive summary section.
        """
        self.doc.add_heading("1. Executive Summary", level=1)
        
        # Overview paragraph
        self.doc.add_paragraph(
            f"{self.company_name} has achieved significant growth and market positioning over the past year, "
            f"with a {self.financial_data['growth_rate']*100:.1f}% increase in revenue and market share of "
            f"{self.financial_data['market_share']*100:.1f}%. Our customer retention rate stands at "
            f"{self.financial_data['customer_retention']*100:.1f}%, reflecting our commitment to customer satisfaction "
            f"and product quality."
        )
        
        # Market position
        self.doc.add_heading("Market Position", level=2)
        self.doc.add_paragraph(
            f"We have strengthened our position in the {fake.bs()} sector, with particular success in "
            f"the {fake.catch_phrase()} segment. Our products continue to receive industry recognition, "
            f"including {random.choice(['awards for innovation', 'customer satisfaction ratings', 'industry certifications'])}."
        )
        
        # Financial highlights
        self.doc.add_heading("Financial Highlights", level=2)
        p = self.doc.add_paragraph("Key financial achievements include:\n")
        financial_highlights = [
            f"Total revenue: ${self.financial_data['revenue']['Total']:,}",
            f"Total profit: ${self.financial_data['profit']['Total']:,}",
            f"Profit margin: {(self.financial_data['profit']['Total'] / self.financial_data['revenue']['Total']) * 100:.1f}%"
        ]
        for highlight in financial_highlights:
            p.add_run(f"• {highlight}\n")
        
        # Strategic direction
        self.doc.add_heading("Strategic Direction", level=2)
        self.doc.add_paragraph(
            f"Our strategy for the coming year focuses on {random.choice(['expanding market share', 'enhancing product offerings', 'entering new markets'])} "
            f"while maintaining our commitment to {random.choice(['customer satisfaction', 'product quality', 'innovation'])}. "
            f"We will continue to invest in {random.choice(['research and development', 'talent acquisition', 'operational efficiency'])} "
            f"to drive sustainable growth."
        )
        
        # Add executive summary chart
        self._add_revenue_chart()
        
        self.doc.add_page_break()
    
    def _add_market_analysis(self):
        """
        Add market analysis section.
        """
        self.doc.add_heading("2. Market Analysis", level=1)
        
        # Industry trends
        self.doc.add_heading("Industry Trends", level=2)
        trends = [
            f"Increasing adoption of {fake.bs()} technologies across industries",
            f"Growing demand for {fake.catch_phrase()} solutions in the {fake.bs()} sector",
            f"Shift towards {random.choice(['cloud-based', 'AI-powered', 'mobile-first'])} platforms",
            f"Rising importance of {random.choice(['data security', 'user experience', 'integration capabilities'])}"
        ]
        p = self.doc.add_paragraph()
        for trend in trends:
            p.add_run(f"• {trend}\n")
        
        # Market size and growth
        self.doc.add_heading("Market Size and Growth", level=2)
        market_size = random.randint(1, 50) * 1000000000
        growth_rate = random.uniform(0.05, 0.2)
        self.doc.add_paragraph(
            f"The global market for {fake.bs()} solutions is valued at approximately ${market_size:,} "
            f"and is expected to grow at a CAGR of {growth_rate*100:.1f}% over the next five years. "
            f"This growth is driven by {random.choice(['increasing digital transformation', 'regulatory changes', 'technological advancements'])} "
            f"across industries."
        )
        
        # Target market segments
        self.doc.add_heading("Target Market Segments", level=2)
        segments = [
            {"name": f"{fake.company_suffix()} Companies", "size": f"${random.randint(100, 999)}M", "growth": f"{random.uniform(0.05, 0.25)*100:.1f}%"},
            {"name": f"{fake.bs()} Providers", "size": f"${random.randint(100, 999)}M", "growth": f"{random.uniform(0.05, 0.25)*100:.1f}%"},
            {"name": f"{fake.catch_phrase()} Sector", "size": f"${random.randint(100, 999)}M", "growth": f"{random.uniform(0.05, 0.25)*100:.1f}%"}
        ]
        
        # Add segments table
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Segment'
        hdr_cells[1].text = 'Market Size'
        hdr_cells[2].text = 'Growth Rate'
        
        for segment in segments:
            row_cells = table.add_row().cells
            row_cells[0].text = segment["name"]
            row_cells[1].text = segment["size"]
            row_cells[2].text = segment["growth"]
        
        self.doc.add_paragraph()
        
        # Geographic distribution
        self.doc.add_heading("Geographic Distribution", level=2)
        self.doc.add_paragraph(
            f"Our primary markets include {random.choice(['North America', 'Europe', 'Asia-Pacific'])}, "
            f"which accounts for {random.randint(40, 80)}% of our revenue. We see significant growth opportunities "
            f"in {random.choice(['emerging markets', 'APAC region', 'Latin America'])}."
        )
        
        self.doc.add_page_break()
    
    def _add_competitor_analysis(self):
        """
        Add competitor analysis section.
        """
        self.doc.add_heading("3. Competitor Analysis", level=1)
        
        # Competitive landscape
        self.doc.add_heading("Competitive Landscape", level=2)
        self.doc.add_paragraph(
            f"The {fake.bs()} market is {random.choice(['highly competitive', 'moderately competitive', 'increasingly competitive'])} "
            f"with several established players and emerging startups. Key competitive factors include "
            f"{random.choice(['product features', 'pricing', 'customer support', 'technological innovation'])}."
        )
        
        # Key competitors
        self.doc.add_heading("Key Competitors", level=2)
        
        for competitor in self.competitors:
            self.doc.add_heading(competitor["name"], level=3)
            p = self.doc.add_paragraph()
            p.add_run(f"Market Share: {competitor['market_share']*100:.1f}%\n")
            
            p.add_run("Strengths:\n")
            for strength in competitor["strengths"]:
                p.add_run(f"• {strength}\n")
            
            p.add_run("Weaknesses:\n")
            for weakness in competitor["weaknesses"]:
                p.add_run(f"• {weakness}\n")
        
        # Competitive advantages
        self.doc.add_heading("Our Competitive Advantages", level=2)
        advantages = [
            f"Superior {fake.bs()} capabilities compared to competitors",
            f"Strong expertise in {fake.catch_phrase()} implementation",
            f"Comprehensive customer support and training programs",
            f"Robust {random.choice(['R&D pipeline', 'partner ecosystem', 'intellectual property portfolio'])}"
        ]
        
        p = self.doc.add_paragraph()
        for advantage in advantages:
            p.add_run(f"• {advantage}\n")
        
        # Add competitor chart
        self._add_competitor_chart()
        
        self.doc.add_page_break()
    
    def _add_strategic_goals(self):
        """
        Add strategic goals section.
        """
        self.doc.add_heading("4. Strategic Goals", level=1)
        
        # Vision and mission
        self.doc.add_heading("Vision and Mission", level=2)
        self.doc.add_paragraph(
            f"Vision: To be the leading provider of {fake.bs()} solutions, transforming how "
            f"{random.choice(['businesses operate', 'customers engage', 'industries evolve'])}.\n\n"
            f"Mission: We deliver innovative {fake.catch_phrase()} solutions that enable our clients "
            f"to {fake.bs()}."
        )
        
        # Strategic objectives
        self.doc.add_heading("Strategic Objectives for Next Fiscal Year", level=2)
        p = self.doc.add_paragraph()
        for i, goal in enumerate(self.strategic_goals):
            p.add_run(f"{i+1}. {goal}\n")
        
        # Key initiatives
        self.doc.add_heading("Key Initiatives", level=2)
        initiatives = [
            {
                "name": f"Expand {fake.bs()} Capabilities",
                "description": f"Enhance our core product with {random.choice(['AI features', 'advanced analytics', 'integration capabilities'])}",
                "timeline": f"Q{random.randint(1, 4)} {self.current_year}",
                "owner": fake.name()
            },
            {
                "name": f"Enter {random.choice(['Asia-Pacific', 'European', 'Latin American'])} Market",
                "description": f"Establish presence in {random.randint(2, 5)} new countries with localized offerings",
                "timeline": f"Q{random.randint(1, 4)} {self.current_year}",
                "owner": fake.name()
            },
            {
                "name": f"Launch {fake.catch_phrase()} Platform",
                "description": f"Introduce new platform targeting the {fake.bs()} vertical",
                "timeline": f"Q{random.randint(1, 4)} {self.current_year}",
                "owner": fake.name()
            }
        ]
        
        # Add initiatives table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Initiative'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Timeline'
        hdr_cells[3].text = 'Owner'
        
        for initiative in initiatives:
            row_cells = table.add_row().cells
            row_cells[0].text = initiative["name"]
            row_cells[1].text = initiative["description"]
            row_cells[2].text = initiative["timeline"]
            row_cells[3].text = initiative["owner"]
        
        self.doc.add_paragraph()
        
        # Success metrics
        self.doc.add_heading("Success Metrics", level=2)
        metrics = [
            f"Revenue growth of {int(self.financial_data['growth_rate']*100)}% year-over-year",
            f"Customer retention rate of {self.financial_data['customer_retention']*100:.1f}%",
            f"Market share increase to {(self.financial_data['market_share']+0.05)*100:.1f}%",
            f"Launch of {random.randint(1, 3)} new products",
            f"Employee satisfaction score above {random.randint(80, 95)}%"
        ]
        
        p = self.doc.add_paragraph()
        for metric in metrics:
            p.add_run(f"• {metric}\n")
        
        self.doc.add_page_break()
    
    def _add_financial_projections(self):
        """
        Add financial projections section.
        """
        self.doc.add_heading("5. Financial Projections", level=1)
        
        # Historical performance
        self.doc.add_heading("Historical Performance", level=2)
        
        # Add financial table
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Q1'
        hdr_cells[2].text = 'Q2'
        hdr_cells[3].text = 'Q3'
        hdr_cells[4].text = 'Q4'
        
        # Add revenue row
        row_cells = table.add_row().cells
        row_cells[0].text = 'Revenue'
        row_cells[1].text = f"${self.financial_data['revenue']['Q1']:,}"
        row_cells[2].text = f"${self.financial_data['revenue']['Q2']:,}"
        row_cells[3].text = f"${self.financial_data['revenue']['Q3']:,}"
        row_cells[4].text = f"${self.financial_data['revenue']['Q4']:,}"
        
        # Add expenses row
        row_cells = table.add_row().cells
        row_cells[0].text = 'Expenses'
        row_cells[1].text = f"${self.financial_data['expenses']['Q1']:,}"
        row_cells[2].text = f"${self.financial_data['expenses']['Q2']:,}"
        row_cells[3].text = f"${self.financial_data['expenses']['Q3']:,}"
        row_cells[4].text = f"${self.financial_data['expenses']['Q4']:,}"
        
        # Add profit row
        row_cells = table.add_row().cells
        row_cells[0].text = 'Profit'
        row_cells[1].text = f"${self.financial_data['profit']['Q1']:,}"
        row_cells[2].text = f"${self.financial_data['profit']['Q2']:,}"
        row_cells[3].text = f"${self.financial_data['profit']['Q3']:,}"
        row_cells[4].text = f"${self.financial_data['profit']['Q4']:,}"
        
        self.doc.add_paragraph()
        
        # Future projections
        self.doc.add_heading("Projections for Next Fiscal Year", level=2)
        self.doc.add_paragraph(
            f"Based on our strategic initiatives and market analysis, we project the following "
            f"financial performance for the next fiscal year:"
        )
        
        projections = [
            f"Revenue: ${self.financial_data['projections']['revenue']:,} "
            f"({((self.financial_data['projections']['revenue'] / self.financial_data['revenue']['Total']) - 1) * 100:.1f}% growth)",
            
            f"Expenses: ${self.financial_data['projections']['expenses']:,} "
            f"({((self.financial_data['projections']['expenses'] / self.financial_data['expenses']['Total']) - 1) * 100:.1f}% increase)",
            
            f"Profit: ${self.financial_data['projections']['profit']:,} "
            f"({((self.financial_data['projections']['profit'] / self.financial_data['profit']['Total']) - 1) * 100:.1f}% growth)",
            
            f"Profit Margin: {(self.financial_data['projections']['profit'] / self.financial_data['projections']['revenue']) * 100:.1f}%"
        ]
        
        p = self.doc.add_paragraph()
        for projection in projections:
            p.add_run(f"• {projection}\n")
        
        # Investment requirements
        self.doc.add_heading("Investment Requirements", level=2)
        investments = [
            {"category": "Product Development", "amount": f"${random.randint(1, 10):,}M", "roi": f"{random.randint(15, 40)}%"},
            {"category": "Market Expansion", "amount": f"${random.randint(1, 5):,}M", "roi": f"{random.randint(15, 40)}%"},
            {"category": "Operational Improvements", "amount": f"${random.randint(1, 3):,}M", "roi": f"{random.randint(15, 40)}%"}
        ]
        
        # Add investments table
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Investment Category'
        hdr_cells[1].text = 'Amount'
        hdr_cells[2].text = 'Expected ROI'
        
        for investment in investments:
            row_cells = table.add_row().cells
            row_cells[0].text = investment["category"]
            row_cells[1].text = investment["amount"]
            row_cells[2].text = investment["roi"]
        
        # Add financial chart
        self._add_projections_chart()
        
        self.doc.add_page_break()
    
    def _add_risk_assessment(self):
        """
        Add risk assessment section.
        """
        self.doc.add_heading("6. Risk Assessment", level=1)
        
        # Risk overview
        self.doc.add_paragraph(
            f"As {self.company_name} pursues its strategic objectives, we have identified several key risks "
            f"that could impact our performance. This section outlines these risks and our mitigation strategies."
        )
        
        # Risk matrix
        self.doc.add_heading("Risk Assessment Matrix", level=2)
        
        # Add risk table
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Risk Category'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Mitigation Strategy'
        
        for risk in self.risks:
            row_cells = table.add_row().cells
            row_cells[0].text = risk["category"]
            row_cells[1].text = risk["description"]
            
            # Generate random mitigation strategy
            mitigation_strategies = [
                f"Implement robust {fake.bs()} procedures",
                f"Diversify {random.choice(['supplier base', 'product portfolio', 'market presence'])}",
                f"Enhance {random.choice(['monitoring systems', 'training programs', 'contingency plans'])}",
                f"Establish strategic partnerships with {fake.company()}",
                f"Invest in {random.choice(['advanced technology', 'redundant systems', 'preventive measures'])}"
            ]
            row_cells[2].text = random.choice(mitigation_strategies)
        
        self.doc.add_paragraph()
        
        # Contingency planning
        self.doc.add_heading("Contingency Planning", level=2)
        self.doc.add_paragraph(
            f"We have developed comprehensive contingency plans for high-impact risks. "
            f"These plans include {random.choice(['alternative supply chains', 'backup systems', 'flexible staffing models'])} "
            f"and are reviewed {random.choice(['quarterly', 'biannually', 'annually'])} by the executive team."
        )
        
        self.doc.add_page_break()
    
    def _add_implementation_plan(self):
        """
        Add implementation plan section.
        """
        self.doc.add_heading("7. Implementation Plan", level=1)
        
        # Timeline
        self.doc.add_heading("Implementation Timeline", level=2)
        quarters = ["Q1", "Q2", "Q3", "Q4"]
        activities = [
            f"Develop {fake.bs()} framework",
            f"Launch {fake.catch_phrase()} initiative",
            f"Expand into {random.choice(['new market', 'new segment', 'new vertical'])}",
            f"Implement {random.choice(['new CRM system', 'ERP upgrade', 'data analytics platform'])}",
            f"Complete {random.choice(['team restructuring', 'process optimization', 'facility expansion'])}"
        ]
        
        # Add timeline table
        table = self.doc.add_table(rows=len(activities)+1, cols=5)
        table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Initiative'
        for i, quarter in enumerate(quarters):
            hdr_cells[i+1].text = quarter
        
        # Add activities
        for i, activity in enumerate(activities):
            row_cells = table.rows[i+1].cells
            row_cells[0].text = activity
            
            # Randomly assign activities to quarters
            start_quarter = random.randint(0, 2)
            duration = random.randint(1, 4 - start_quarter)
            
            for q in range(start_quarter, start_quarter + duration):
                if q < 4:  # Ensure we don't go beyond Q4
                    row_cells[q+1].text = '✓'
        
        self.doc.add_paragraph()
        
        # Resource allocation
        self.doc.add_heading("Resource Allocation", level=2)
        resources = [
            {"department": "Product Development", "headcount": f"{random.randint(10, 50)}", "budget": f"${random.randint(1, 10):,}M"},
            {"department": "Sales & Marketing", "headcount": f"{random.randint(10, 50)}", "budget": f"${random.randint(1, 10):,}M"},
            {"department": "Operations", "headcount": f"{random.randint(10, 50)}", "budget": f"${random.randint(1, 5):,}M"},
            {"department": "Customer Support", "headcount": f"{random.randint(10, 30)}", "budget": f"${random.randint(1, 3):,}M"}
        ]
        
        # Add resources table
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Department'
        hdr_cells[1].text = 'Headcount'
        hdr_cells[2].text = 'Budget'
        
        for resource in resources:
            row_cells = table.add_row().cells
            row_cells[0].text = resource["department"]
            row_cells[1].text = resource["headcount"]
            row_cells[2].text = resource["budget"]
        
        self.doc.add_paragraph()
        
        # Key milestones
        self.doc.add_heading("Key Milestones", level=2)
        milestones = [
            f"{quarters[0]}: Complete {fake.bs()} strategy",
            f"{quarters[1]}: Launch {fake.catch_phrase()} product",
            f"{quarters[2]}: Achieve {random.randint(100, 500)} {random.choice(['customers', 'users', 'installations'])}",
            f"{quarters[3]}: Reach ${random.randint(1, 10)}M {random.choice(['revenue', 'sales', 'bookings'])}"
        ]
        
        p = self.doc.add_paragraph()
        for milestone in milestones:
            p.add_run(f"• {milestone}\n")
        
        # Governance
        self.doc.add_heading("Governance Structure", level=2)
        self.doc.add_paragraph(
            f"Our implementation will be overseen by a steering committee comprising executives from "
            f"{random.choice(['Product', 'Sales', 'Marketing', 'Operations', 'Finance'])} and "
            f"{random.choice(['Product', 'Sales', 'Marketing', 'Operations', 'Finance'])}. "
            f"The committee will meet {random.choice(['weekly', 'biweekly', 'monthly'])} to review progress, "
            f"address issues, and ensure alignment with strategic objectives."
        )
        
        self.doc.add_page_break()
    
    def _add_appendices(self):
        """
        Add appendices section.
        """
        self.doc.add_heading("8. Appendices", level=1)
        
        # Team bios
        self.doc.add_heading("Appendix A: Executive Team Biographies", level=2)
        executives = [
            {"name": fake.name(), "title": "Chief Executive Officer", "bio": fake.paragraph(nb_sentences=3)},
            {"name": fake.name(), "title": "Chief Financial Officer", "bio": fake.paragraph(nb_sentences=3)},
            {"name": fake.name(), "title": "Chief Technology Officer", "bio": fake.paragraph(nb_sentences=3)},
            {"name": fake.name(), "title": "Chief Marketing Officer", "bio": fake.paragraph(nb_sentences=3)}
        ]
        
        for executive in executives:
            self.doc.add_heading(f"{executive['name']}, {executive['title']}", level=3)
            self.doc.add_paragraph(executive["bio"])
        
        # Market research
        self.doc.add_heading("Appendix B: Detailed Market Research", level=2)
        self.doc.add_paragraph(
            f"This appendix contains additional market research data supporting our strategic plans. "
            f"The research was conducted by {fake.company()} in {fake.date_this_year().strftime('%B %Y')} "
            f"and included surveys of {random.randint(100, 1000)} potential customers in "
            f"{random.randint(3, 10)} countries."
        )
        
        # Add some sample data
        self.doc.add_heading("Customer Survey Results", level=3)
        survey_results = [
            {"question": "Importance of product features", "response": f"{random.randint(60, 90)}% rated as 'Very Important'"},
            {"question": "Price sensitivity", "response": f"{random.randint(40, 70)}% would pay premium for quality"},
            {"question": "Brand awareness", "response": f"{random.randint(30, 80)}% familiar with our brand"}
        ]
        
        # Add survey table
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Survey Question'
        hdr_cells[1].text = 'Response'
        
        for result in survey_results:
            row_cells = table.add_row().cells
            row_cells[0].text = result["question"]
            row_cells[1].text = result["response"]
    
    def _add_revenue_chart(self):
        """
        Add revenue chart to the document.
        """
        # Create chart data
        quarters = ["Q1", "Q2", "Q3", "Q4"]
        revenue_data = [
            self.financial_data["revenue"]["Q1"] / 1000000,
            self.financial_data["revenue"]["Q2"] / 1000000,
            self.financial_data["revenue"]["Q3"] / 1000000,
            self.financial_data["revenue"]["Q4"] / 1000000
        ]
        profit_data = [
            self.financial_data["profit"]["Q1"] / 1000000,
            self.financial_data["profit"]["Q2"] / 1000000,
            self.financial_data["profit"]["Q3"] / 1000000,
            self.financial_data["profit"]["Q4"] / 1000000
        ]
        
        # Create the chart
        plt.figure(figsize=(10, 5))
        x = np.arange(len(quarters))
        width = 0.35
        
        plt.bar(x - width/2, revenue_data, width, label='Revenue')
        plt.bar(x + width/2, profit_data, width, label='Profit')
        
        plt.xlabel('Quarter')
        plt.ylabel('Millions ($)')
        plt.title('Quarterly Revenue and Profit')
        plt.xticks(x, quarters)
        plt.legend()
        
        # Save chart
        chart_path = os.path.join(self.output_dir, "revenue_chart.png")
        plt.savefig(chart_path)
        plt.close()
        
        # Add chart to document
        self.doc.add_paragraph().add_run().add_break()
        self.doc.add_picture(chart_path, width=Inches(6))
        
        # Add caption
        caption = self.doc.add_paragraph("Figure 1: Quarterly Revenue and Profit")
        caption.alignment = 1  # Center
    
    def _add_competitor_chart(self):
        """
        Add competitor market share chart.
        """
        # Create chart data
        companies = [self.company_name] + [comp["name"] for comp in self.competitors]
        market_shares = [self.financial_data["market_share"]] + [comp["market_share"] for comp in self.competitors]
        
        # Create the chart
        plt.figure(figsize=(10, 5))
        plt.pie(market_shares, labels=None, autopct='%1.1f%%', shadow=True, startangle=90)
        plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle
        plt.title('Market Share Distribution')
        
        # Add legend
        plt.legend(companies, loc="best")
        
        # Save chart
        chart_path = os.path.join(self.output_dir, "market_share_chart.png")
        plt.savefig(chart_path)
        plt.close()
        
        # Add chart to document
        self.doc.add_paragraph().add_run().add_break()
        self.doc.add_picture(chart_path, width=Inches(6))
        
        # Add caption
        caption = self.doc.add_paragraph("Figure 2: Market Share Distribution")
        caption.alignment = 1  # Center
    
    def _add_projections_chart(self):
        """
        Add financial projections chart.
        """
        # Create chart data
        categories = ['Revenue', 'Expenses', 'Profit']
        current_year = [
            self.financial_data["revenue"]["Total"] / 1000000,
            self.financial_data["expenses"]["Total"] / 1000000,
            self.financial_data["profit"]["Total"] / 1000000
        ]
        next_year = [
            self.financial_data["projections"]["revenue"] / 1000000,
            self.financial_data["projections"]["expenses"] / 1000000,
            self.financial_data["projections"]["profit"] / 1000000
        ]
        
        # Create the chart
        plt.figure(figsize=(10, 5))
        x = np.arange(len(categories))
        width = 0.35
        
        plt.bar(x - width/2, current_year, width, label=f'Current Year ({self.current_year})')
        plt.bar(x + width/2, next_year, width, label=f'Next Year ({self.current_year+1})')
        
        plt.xlabel('Category')
        plt.ylabel('Millions ($)')
        plt.title('Financial Projections')
        plt.xticks(x, categories)
        plt.legend()
        
        # Save chart
        chart_path = os.path.join(self.output_dir, "projections_chart.png")
        plt.savefig(chart_path)
        plt.close()
        
        # Add chart to document
        self.doc.add_paragraph().add_run().add_break()
        self.doc.add_picture(chart_path, width=Inches(6))
        
        # Add caption
        caption = self.doc.add_paragraph("Figure 3: Financial Projections Comparison")
        caption.alignment = 1  # Center
    
    def _convert_to_pdf(self, docx_path):
        """
        Convert the document to PDF if possible.
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            Path to the PDF file if conversion successful, None otherwise
        """
        try:
            # Check if docx2pdf is available
            from docx2pdf import convert
            
            # Convert to PDF
            pdf_path = docx_path.replace(".docx", ".pdf")
            convert(docx_path, pdf_path)
            
            return pdf_path
        except ImportError:
            print("docx2pdf not installed, skipping PDF conversion")
            return None
        except Exception as e:
            print(f"Error converting to PDF: {e}")
            return None


def main():
    """
    Generate a sample business document.
    """
    # Create output directory
    output_dir = "samples"
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate the document
    generator = SampleDocGenerator(output_dir=output_dir)
    document_path = generator.generate_document()
    
    print(f"Sample document generated: {document_path}")


if __name__ == "__main__":
    main()
